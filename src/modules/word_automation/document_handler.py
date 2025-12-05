"""
Document Handler for Word automation.

Provides comprehensive Word document manipulation using python-docx.
"""
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class DocumentHandler:
    """
    Handle Word document operations.

    Supports:
    - Creating new documents
    - Loading existing templates
    - Adding headings, paragraphs, tables
    - Replacing placeholders
    - Applying formatting
    - Extracting document structure
    """

    def __init__(self, filepath: Optional[str] = None) -> None:
        """
        Initialize document handler.

        Args:
            filepath: Optional path to existing document to load
        """
        self.logger = logging.getLogger(__name__)
        self.filepath = filepath
        self.document = None

        if filepath:
            self.load_template(filepath)
        else:
            self.logger.info("Document handler initialized (no document loaded)")

    def create_document(self) -> 'DocumentHandler':
        """Create a new blank document."""
        self.document = Document()
        self.filepath = None
        self.logger.info("Created new blank document")
        return self

    def load_template(self, filepath: str) -> 'DocumentHandler':
        """
        Load an existing Word document as template.

        Args:
            filepath: Path to .docx file

        Returns:
            self for method chaining

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid .docx
        """
        filepath = Path(filepath)

        if not filepath.exists():
            raise FileNotFoundError(f"Document not found: {filepath}")

        if filepath.suffix.lower() != '.docx':
            raise ValueError(f"File must be .docx format: {filepath}")

        try:
            self.document = Document(str(filepath))
            self.filepath = str(filepath)
            self.logger.info(f"Loaded document template: {filepath}")
            return self
        except Exception as e:
            self.logger.error(f"Failed to load document: {e}")
            raise

    def add_heading(self, text: str, level: int = 1) -> 'DocumentHandler':
        """
        Add a heading to the document.

        Args:
            text: Heading text
            level: Heading level (1-9, where 1 is largest)

        Returns:
            self for method chaining
        """
        if not self.document:
            self.create_document()

        level = max(1, min(9, level))  # Clamp to 1-9
        self.document.add_heading(text, level=level)
        self.logger.debug(f"Added heading (level {level}): {text}")
        return self

    def add_paragraph(self, text: str, style: Optional[str] = None) -> 'DocumentHandler':
        """
        Add a paragraph to the document.

        Args:
            text: Paragraph text
            style: Optional style name ('Normal', 'BodyText', 'Quote', etc.)

        Returns:
            self for method chaining
        """
        if not self.document:
            self.create_document()

        if style:
            para = self.document.add_paragraph(text, style=style)
        else:
            para = self.document.add_paragraph(text)

        self.logger.debug(f"Added paragraph: {text[:50]}...")
        return self

    def add_page_break(self) -> 'DocumentHandler':
        """Add a page break."""
        if not self.document:
            self.create_document()

        self.document.add_page_break()
        self.logger.debug("Added page break")
        return self

    def add_table(self, data: List[List[Any]], headers: Optional[List[str]] = None, style: str = 'Light Grid Accent 1') -> 'DocumentHandler':
        """
        Add a table to the document.

        Args:
            data: 2D list of table data (rows and columns)
            headers: Optional list of column headers
            style: Table style name

        Returns:
            self for method chaining

        Example:
            handler.add_table(
                data=[['John', 25], ['Jane', 30]],
                headers=['Name', 'Age']
            )
        """
        if not self.document:
            self.create_document()

        # Calculate table dimensions
        num_cols = len(headers) if headers else len(data[0]) if data else 0
        num_rows = (1 if headers else 0) + len(data)

        if num_cols == 0 or num_rows == 0:
            self.logger.warning("Cannot create table with no data")
            return self

        # Create table
        table = self.document.add_table(rows=num_rows, cols=num_cols)

        # Apply style
        if style:
            try:
                table.style = style
            except KeyError:
                self.logger.warning(f"Table style not found: {style}, using default")

        # Add headers
        row_idx = 0
        if headers:
            header_row = table.rows[0]
            for col_idx, header_text in enumerate(headers):
                header_row.cells[col_idx].text = str(header_text)
                # Make headers bold
                for paragraph in header_row.cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            row_idx = 1

        # Add data
        for data_row in data:
            table_row = table.rows[row_idx]
            for col_idx, cell_value in enumerate(data_row):
                table_row.cells[col_idx].text = str(cell_value)
            row_idx += 1

        self.logger.debug(f"Added table with {num_rows} rows and {num_cols} columns")
        return self

    def replace_placeholders(self, mapping: Dict[str, Any], placeholder_pattern: str = r'\{\{(\w+)\}\}') -> 'DocumentHandler':
        """
        Replace placeholders in the document with actual values.

        Searches through all paragraphs and table cells for placeholders
        matching the pattern and replaces them with values from the mapping.

        Args:
            mapping: Dictionary mapping placeholder names to values
            placeholder_pattern: Regex pattern for placeholders (default: {{variable}})

        Returns:
            self for method chaining

        Example:
            handler.replace_placeholders({
                'client_name': 'Acme Corp',
                'date': '2025-11-04',
                'amount': '$1,500.00'
            })
        """
        if not self.document:
            self.logger.error("No document loaded")
            return self

        replacements_made = 0

        # Replace in paragraphs
        for paragraph in self.document.paragraphs:
            for match in re.finditer(placeholder_pattern, paragraph.text):
                placeholder_name = match.group(1)
                if placeholder_name in mapping:
                    # Replace in the paragraph text
                    old_text = paragraph.text
                    new_text = old_text.replace(match.group(0), str(mapping[placeholder_name]))

                    # Clear existing runs and add new text
                    # This preserves formatting better than direct replacement
                    paragraph.text = new_text
                    replacements_made += 1

        # Replace in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for match in re.finditer(placeholder_pattern, paragraph.text):
                            placeholder_name = match.group(1)
                            if placeholder_name in mapping:
                                old_text = paragraph.text
                                new_text = old_text.replace(match.group(0), str(mapping[placeholder_name]))
                                paragraph.text = new_text
                                replacements_made += 1

        self.logger.info(f"Replaced {replacements_made} placeholders")
        return self

    def apply_formatting(self, paragraph_index: int = -1, **formatting: Any) -> 'DocumentHandler':
        """
        Apply formatting to a paragraph or the last added paragraph.

        Args:
            paragraph_index: Index of paragraph to format (-1 for last)
            **formatting: Formatting options:
                - font_size: Font size in points
                - font_name: Font family name
                - bold: True/False
                - italic: True/False
                - color: RGB tuple (r, g, b)
                - alignment: 'left', 'center', 'right', 'justify'

        Returns:
            self for method chaining
        """
        if not self.document or not self.document.paragraphs:
            self.logger.warning("No paragraphs to format")
            return self

        para = self.document.paragraphs[paragraph_index]

        # Apply paragraph-level formatting
        if 'alignment' in formatting:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            para.alignment = alignment_map.get(formatting['alignment'], WD_ALIGN_PARAGRAPH.LEFT)

        # Apply run-level formatting to all runs
        for run in para.runs:
            if 'font_size' in formatting:
                run.font.size = Pt(formatting['font_size'])

            if 'font_name' in formatting:
                run.font.name = formatting['font_name']

            if 'bold' in formatting:
                run.font.bold = formatting['bold']

            if 'italic' in formatting:
                run.font.italic = formatting['italic']

            if 'color' in formatting:
                r, g, b = formatting['color']
                run.font.color.rgb = RGBColor(r, g, b)

        self.logger.debug(f"Applied formatting to paragraph {paragraph_index}")
        return self

    def extract_structure(self) -> Dict[str, Any]:
        """
        Extract the structure and content of the document.

        Returns a dictionary containing:
        - headings: List of headings with level and text
        - paragraphs: List of paragraph texts
        - tables: List of tables with headers and data
        - placeholders: Set of detected placeholder names

        Returns:
            Dictionary with document structure
        """
        if not self.document:
            self.logger.error("No document loaded")
            return {}

        structure = {
            'headings': [],
            'paragraphs': [],
            'tables': [],
            'placeholders': set(),
            'total_paragraphs': len(self.document.paragraphs),
            'total_tables': len(self.document.tables)
        }

        # Extract headings and paragraphs
        for para in self.document.paragraphs:
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                structure['headings'].append({
                    'level': level,
                    'text': para.text
                })

            # Add all paragraphs
            if para.text.strip():
                structure['paragraphs'].append(para.text)

            # Find placeholders
            placeholders = re.findall(r'\{\{(\w+)\}\}', para.text)
            structure['placeholders'].update(placeholders)

        # Extract tables
        for table_idx, table in enumerate(self.document.tables):
            table_data = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': []
            }

            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data['data'].append(row_data)

                # Check for placeholders in table cells
                for cell_text in row_data:
                    placeholders = re.findall(r'\{\{(\w+)\}\}', cell_text)
                    structure['placeholders'].update(placeholders)

            structure['tables'].append(table_data)

        # Convert set to sorted list for consistency
        structure['placeholders'] = sorted(list(structure['placeholders']))

        self.logger.info(f"Extracted structure: {len(structure['paragraphs'])} paragraphs, "
                        f"{len(structure['tables'])} tables, {len(structure['placeholders'])} placeholders")

        return structure

    def save(self, filepath: Optional[str] = None) -> str:
        """
        Save the document to a file.

        Args:
            filepath: Output path for .docx file (uses self.filepath if not provided)

        Raises:
            ValueError: If no filepath provided and no filepath set
        """
        if not self.document:
            raise ValueError("No document to save")

        output_path = filepath if filepath else self.filepath

        if not output_path:
            raise ValueError("No filepath provided for saving document")

        # Ensure .docx extension
        output_path = Path(output_path)
        if output_path.suffix.lower() != '.docx':
            output_path = output_path.with_suffix('.docx')

        # Create parent directory if needed
        output_path.parent.mkdir(parents=True, exist_ok=True)

        self.document.save(str(output_path))
        self.filepath = str(output_path)
        self.logger.info(f"Saved document to: {output_path}")

        return str(output_path)

    def get_document(self) -> Document:
        """
        Get the underlying python-docx Document object.

        Returns:
            Document object or None if not created
        """
        return self.document
